"""
访谈逐字稿清洗服务

流程：读取原始 docx → 应用规则清洗 → Claude API 语义处理 → 输出格式化 docx
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import anthropic


SPEAKER_COLORS = {
    1: RGBColor(0xFF, 0x00, 0x00),  # 访谈者 - 红色
    2: RGBColor(0x00, 0x00, 0x00),  # 受访者 - 黑色
    3: RGBColor(0x00, 0x00, 0xFF),  # 第三人 - 蓝色
}

SPEAKER_PATTERN = re.compile(r"^发言人\s*\d+\s*$")

PURE_RESPONSE_PATTERN = re.compile(
    r"^(嗯|哦|好的|对对对|对的|好|是的|嗯嗯|哦哦|对|是|OK|ok|嗯嗯嗯|好好好|对对|啊|哈|嗯哼|没错)[。，、.!！？]*$"
)

FILLER_PATTERNS = [
    (re.compile(r"呃[，,\s]*"), ""),
]

STACKING_PATTERNS = [
    (re.compile(r"然后还有就是"), "然后还有"),
    (re.compile(r"比如看一些像"), "像"),
    (re.compile(r"去那个去"), "去"),
    (re.compile(r"就是那个就是"), "就是"),
    (re.compile(r"然后就是然后"), "然后"),
]

EXTRA_SPACE_PATTERN = re.compile(r"(?<=[一-鿿])\s+(?=[一-鿿\w])")


def _identify_speakers(paragraphs: list[str]) -> list[dict]:
    """Parse raw paragraphs into structured segments with speaker identification."""
    segments = []
    current_speaker = 1
    current_text_parts = []

    for para in paragraphs:
        text = para.strip()
        if not text:
            continue

        match = re.match(r"^发言人\s*(\d+)\s*$", text)
        if match:
            if current_text_parts:
                segments.append({
                    "speaker": current_speaker,
                    "text": "\n".join(current_text_parts),
                })
                current_text_parts = []
            current_speaker = int(match.group(1))
            continue

        current_text_parts.append(text)

    if current_text_parts:
        segments.append({
            "speaker": current_speaker,
            "text": "\n".join(current_text_parts),
        })

    return segments


def _is_metadata(text: str) -> bool:
    """Check if a line is auto-generated metadata."""
    metadata_indicators = [
        "创建时间", "转写时长", "关键词", "时长：", "字数：",
        "转录引擎", "文件大小", "音频时长", "开始时间", "结束时间",
    ]
    return any(ind in text for ind in metadata_indicators)


def _remove_pure_responses(segments: list[dict]) -> list[dict]:
    """Remove segments that are pure acknowledgment responses."""
    result = []
    for seg in segments:
        lines = seg["text"].split("\n")
        cleaned_lines = [
            line for line in lines
            if not PURE_RESPONSE_PATTERN.match(line.strip())
        ]
        if cleaned_lines:
            seg["text"] = "\n".join(cleaned_lines)
            result.append(seg)
    return result


def _apply_regex_rules(text: str) -> str:
    """Apply filler removal, stacking cleanup, and space normalization."""
    for pattern, replacement in FILLER_PATTERNS:
        text = pattern.sub(replacement, text)

    for pattern, replacement in STACKING_PATTERNS:
        text = pattern.sub(replacement, text)

    text = EXTRA_SPACE_PATTERN.sub("", text)

    return text


def _merge_same_speaker(segments: list[dict]) -> list[dict]:
    """Merge consecutive segments from the same speaker."""
    if not segments:
        return []
    merged = [segments[0].copy()]
    for seg in segments[1:]:
        if seg["speaker"] == merged[-1]["speaker"]:
            merged[-1]["text"] += "\n" + seg["text"]
        else:
            merged.append(seg.copy())
    return merged


def _call_claude_for_semantic_processing(segments: list[dict], api_key: str, base_url: str = "", model: str = "claude-sonnet-4-6") -> list[dict]:
    """
    Use Claude API to:
    1. Identify and remove abandoned restart fragments (rule 3.5)
    2. Insert topic headings at topic transitions (rule 3.9)
    """
    client_kwargs = {"api_key": api_key}
    if base_url:
        client_kwargs["base_url"] = base_url
    client = anthropic.Anthropic(**client_kwargs)

    full_text = ""
    for i, seg in enumerate(segments):
        speaker_label = "访谈者" if seg["speaker"] == 1 else f"发言人{seg['speaker']}"
        full_text += f"[段落{i}][{speaker_label}]\n{seg['text']}\n\n"

    prompt = f"""你是一个访谈逐字稿清洗助手。请对以下逐字稿执行两个任务：

## 任务1：识别并删除废弃重启片段
说话人说到一半放弃、重新组织语言的部分，只保留重启后的最终表达。
判断依据：同一句话中出现"就是...就是..."或类似重复连接词引导的修正结构，前半截为废弃片段。
示例："就是像一些客，就是客户的一些留言"→"客户的一些留言"

## 任务2：识别话题转换点并插入小标题
识别访谈者（发言人1）转换话题的位置。在该段落前应该插入一个话题小标题。
判断依据：访谈者用明确的过渡语切换到新的讨论主题。
小标题命名参考维度：基本信息、购买旅程、旧方案、使用场景-XX、使用过程感受、评价、需求、替代方案-XX、日常使用习惯等（根据实际内容灵活命名）。

## 输出格式
严格按JSON格式返回，不要其他文字：
{{
  "cleaned_segments": [
    {{"index": 0, "text": "清洗后的文本", "heading_before": null}},
    {{"index": 1, "text": "清洗后的文本", "heading_before": "基本信息"}},
    ...
  ]
}}

每个segment的index对应输入段落编号，text是清洗后的内容（只做废弃片段删除，不改写其他内容），heading_before是该段落前需要插入的小标题（null表示不需要）。

## 逐字稿内容：
{full_text}"""

    response = client.messages.create(
        model=model,
        max_tokens=8000,
        messages=[{"role": "user", "content": prompt}],
    )

    response_text = response.content[0].text

    import json
    json_match = re.search(r"\{[\s\S]*\}", response_text)
    if not json_match:
        return segments

    try:
        result = json.loads(json_match.group())
    except json.JSONDecodeError:
        return segments

    cleaned = result.get("cleaned_segments", [])
    processed_segments = []
    for item in cleaned:
        idx = item.get("index", 0)
        if idx < len(segments):
            seg = segments[idx].copy()
            seg["text"] = item.get("text", seg["text"])
            seg["heading_before"] = item.get("heading_before")
            processed_segments.append(seg)

    return processed_segments if processed_segments else segments


def _build_output_document(segments: list[dict], title: str) -> Document:
    """Build the final formatted docx document."""
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x36, 0x5F, 0x91)

    for seg in segments:
        heading_text = seg.get("heading_before")
        if heading_text:
            h_para = doc.add_paragraph()
            h_run = h_para.add_run(heading_text)
            h_run.bold = True
            h_run.font.size = Pt(12)
            h_para.paragraph_format.space_before = Pt(12)
            h_para.paragraph_format.space_after = Pt(6)

        color = SPEAKER_COLORS.get(seg["speaker"], RGBColor(0x00, 0x00, 0x00))
        for line in seg["text"].split("\n"):
            if line.strip():
                para = doc.add_paragraph()
                run = para.add_run(line.strip())
                run.font.size = Pt(11)
                run.font.color.rgb = color

    return doc


def clean_transcript(input_path: str, output_dir: str, api_key: str, base_url: str = "", model: str = "claude-sonnet-4-6") -> str:
    """
    Main entry point: clean a raw interview transcript docx file.

    Args:
        input_path: Path to the raw transcript .docx file
        output_dir: Directory to save the cleaned file
        api_key: Anthropic API key for Claude calls
        base_url: Custom Anthropic API base URL (e.g. CloudFront proxy)
        model: Model name to use for semantic processing

    Returns:
        Path to the cleaned output file
    """
    doc = Document(input_path)

    raw_paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not _is_metadata(text):
            raw_paragraphs.append(text)

    segments = _identify_speakers(raw_paragraphs)
    segments = _remove_pure_responses(segments)

    for seg in segments:
        seg["text"] = _apply_regex_rules(seg["text"])

    segments = _merge_same_speaker(segments)
    segments = [s for s in segments if s["text"].strip()]

    segments = _call_claude_for_semantic_processing(segments, api_key, base_url, model)

    input_name = Path(input_path).stem
    title = f"{input_name}（清洗版）"

    output_doc = _build_output_document(segments, title)

    output_filename = f"{input_name}（清洗后版本）.docx"
    output_path = os.path.join(output_dir, output_filename)
    output_doc.save(output_path)

    return output_path
